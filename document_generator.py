import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from app import app

def generate_der_document(client):
    """Génère un Document d'Entrée en Relation (DER) pour le client"""
    try:
        # Créer un nouveau document
        doc = Document()
        
        # En-tête du conseiller
        header_para = doc.add_paragraph()
        header_para.add_run('Kim BUI, en entreprise individuelle (micro-entreprise)\n').bold = True
        header_para.add_run('N° SIREN : 930 610 084 – APE : 68.31Z\n')
        header_para.add_run('Conseiller en Investissements Financiers (CIF), membre de l\'ANACOFI-CIF, association agréée par l\'AMF\n')
        header_para.add_run('Immatriculé à l\'ORIAS n° 24006088 (www.orias.fr)\n')
        header_para.add_run('Assurance Responsabilité Civile Professionnelle : ZURICH INSURANCE PLC – Police n° 7400023129\n')
        header_para.add_run('Siège social : 105 rue du 4 août 1789 – 69100 Villeurbanne\n')
        header_para.add_run('📧 kim.bui.cgp@gmail.com – ☎ 06.01.33.10.60\n\n')
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Titre
        title = doc.add_heading('DOCUMENT D\'ENTRÉE EN RELATION (DER)', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph('\n')
        
        # Date d'entrée en relation
        if client.date_entree_relation:
            date_para = doc.add_paragraph()
            date_para.add_run('Date d\'entrée en relation : ').bold = True
            date_para.add_run(client.date_entree_relation.strftime('%d/%m/%Y'))
        
        doc.add_paragraph('\n')
        
        # Informations client
        doc.add_heading('IDENTIFICATION DU CLIENT', level=1)
        
        client_info = doc.add_paragraph()
        client_info.add_run('Nom et Prénom : ').bold = True
        client_info.add_run(f'{client.prenom} {client.nom}\n')
        
        client_info.add_run('Adresse électronique : ').bold = True
        client_info.add_run(f'{client.email}\n')
        
        if client.telephone:
            client_info.add_run('Téléphone : ').bold = True
            client_info.add_run(f'{client.telephone}\n')
        
        if client.date_naissance:
            client_info.add_run('Date de naissance : ').bold = True
            client_info.add_run(f'{client.date_naissance.strftime("%d/%m/%Y")}\n')
        
        if client.adresse:
            client_info.add_run('Adresse : ').bold = True
            client_info.add_run(f'{client.adresse}\n')
        
        if client.profession:
            client_info.add_run('Profession : ').bold = True
            client_info.add_run(f'{client.profession}\n')
        
        # Ville du client (demandée par l'utilisateur)
        if hasattr(client, 'ville') and client.ville:
            client_info.add_run('Ville : ').bold = True
            client_info.add_run(f'{client.ville}\n')
        
        # STATUTS LÉGAUX ET AUTORITÉS DE TUTELLE
        doc.add_heading('STATUTS LÉGAUX ET AUTORITÉS DE TUTELLE', level=1)
        
        statuts_text = """Votre conseiller est immatriculé au Registre Unique des Intermédiaires en Assurance, Banque et Finance (ORIAS) sous le n° 24006088 pour les activités réglementées suivantes:

CIF enregistré auprès de l'ANACOFI-CIF, contrôlable par l'AMF.

Couverture en Responsabilité Civile Professionnelle et Garantie Financière souscrites auprès de ZURICH INSURANCE PLC, n°7400023129.

Montants:
• Responsabilité Civile Professionnelle: 1 000 000,0 €
• Garantie financière: Non approprié, sauf exception

Engagement à respecter le Code de Bonne Conduite de l'ANACOFI-CIF.

Notre cabinet ne prend pas en compte les facteurs de durabilité dans la sélection des instruments financiers proposés.

Liste des principaux partenaires:
Shares Financial Assets, entreprise d'investissement agréée sous le numéro CIB 17183 par l'Autorité de Contrôle Prudentiel et de Résolution, avec une convention de distribution de produits financiers et rémunération par rétrocessions/commissions."""
        
        doc.add_paragraph(statuts_text)
        
        # MODE DE FACTURATION ET RÉMUNÉRATION
        doc.add_heading('MODE DE FACTURATION ET RÉMUNÉRATION DU PROFESSIONNEL EN CIF', level=1)
        
        remuneration_text = """Bilan patrimonial obligatoire avant toute recommandation ou allocation d'actifs, avec tarifs selon complexité.

Accompagnement et conseil personnalisé avec honoraires calculés sur encours et/ou rétrocessions selon produits.

Coaching patrimonial à 200€ TTC par session d'une heure.

Modalités générales:
• Conseil non-indépendant
• Communication transparente des modes de rémunération
• Devis préalable pour prestations spécifiques

Mode de communication: Email, téléphone, visio, SMS.

Traitement des réclamations: Engagement à traiter dans les délais, possibilité de saisir un médiateur compétent."""
        
        doc.add_paragraph(remuneration_text)
        
        # Documents à fournir
        doc.add_heading('DOCUMENTS À FOURNIR', level=1)
        
        docs_text = """Pour poursuivre cette relation et vous fournir un conseil adapté, vous devrez nous transmettre :

Documents obligatoires :
• Pièce d'identité en cours de validité
• Avis d'imposition de l'année en cours
• Justificatif de domicile de moins de 3 mois
• Relevés de comptes bancaires des 3 derniers mois

Ces documents nous permettront d'évaluer votre situation financière et de vous proposer des solutions adaptées."""
        
        doc.add_paragraph(docs_text)
        
        # DATE ET SIGNATURE
        doc.add_heading('DATE ET SIGNATURE', level=1)
        
        doc.add_paragraph('\n')
        signature_table = doc.add_table(rows=4, cols=2)
        signature_table.style = 'Table Grid'
        
        # En-têtes
        signature_table.cell(0, 0).text = 'Le client'
        signature_table.cell(0, 1).text = 'Le conseiller'
        
        # Lieux
        ville_client = getattr(client, 'ville', '[ville du client]') if hasattr(client, 'ville') else '[ville du client]'
        signature_table.cell(1, 0).text = f'Fait à : {ville_client}'
        signature_table.cell(1, 1).text = 'Fait à : Villeurbanne'
        
        # Dates
        signature_table.cell(2, 0).text = 'Date et signature :'
        signature_table.cell(2, 1).text = 'Date : 08/07/2025'
        
        # Espaces pour signatures
        signature_table.cell(3, 0).text = '\n\n\n'
        signature_table.cell(3, 1).text = 'Signature :\n\n\n'
        
        # Sauvegarder le document
        filename = f'der_{client.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        filepath = os.path.join(app.config['GENERATED_DOCS_FOLDER'], filename)
        doc.save(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Erreur génération DER: {e}")
        return None

def generate_investment_report(client):
    """Génère un rapport d'adéquation pour le client"""
    try:
        # Créer un nouveau document
        doc = Document()
        
        # Titre
        title = doc.add_heading('RAPPORT D\'ADÉQUATION', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Informations client
        doc.add_heading('INFORMATIONS CLIENT', level=1)
        
        client_info = doc.add_paragraph()
        client_info.add_run('Nom: ').bold = True
        client_info.add_run(f'{client.prenom} {client.nom}\n')
        
        client_info.add_run('Email: ').bold = True
        client_info.add_run(f'{client.email}\n')
        
        if client.telephone:
            client_info.add_run('Téléphone: ').bold = True
            client_info.add_run(f'{client.telephone}\n')
        
        if client.date_naissance:
            client_info.add_run('Date de naissance: ').bold = True
            client_info.add_run(f'{client.date_naissance.strftime("%d/%m/%Y")}\n')
        
        if client.profession:
            client_info.add_run('Profession: ').bold = True
            client_info.add_run(f'{client.profession}\n')
        
        # Situation financière
        doc.add_heading('SITUATION FINANCIÈRE', level=1)
        
        finance_info = doc.add_paragraph()
        if client.revenus_mensuels:
            finance_info.add_run('Revenus mensuels: ').bold = True
            finance_info.add_run(f'{client.revenus_mensuels:,.0f} €\n')
        
        if client.patrimoine_total:
            finance_info.add_run('Patrimoine total: ').bold = True
            finance_info.add_run(f'{client.patrimoine_total:,.0f} €\n')
        
        if client.charges_mensuelles:
            finance_info.add_run('Charges mensuelles: ').bold = True
            finance_info.add_run(f'{client.charges_mensuelles:,.0f} €\n')
        
        # Profil investisseur
        doc.add_heading('PROFIL INVESTISSEUR', level=1)
        
        profil_info = doc.add_paragraph()
        if client.tolerance_risque:
            profil_info.add_run('Tolérance au risque: ').bold = True
            profil_info.add_run(f'{client.tolerance_risque.value}\n')
        
        if client.horizon_investissement:
            profil_info.add_run('Horizon d\'investissement: ').bold = True
            profil_info.add_run(f'{client.horizon_investissement.value}\n')
        
        if client.experience_financiere:
            profil_info.add_run('Expérience financière: ').bold = True
            profil_info.add_run(f'{client.experience_financiere}\n')
        
        if client.profil_score:
            profil_info.add_run('Score de profil: ').bold = True
            profil_info.add_run(f'{client.profil_score}/5\n')
        
        # Recommandations
        doc.add_heading('RECOMMANDATIONS', level=1)
        
        recommandations = get_recommandations(client)
        for rec in recommandations:
            doc.add_paragraph(rec, style='List Bullet')
        
        # Date et signature
        doc.add_paragraph('\n\n')
        signature = doc.add_paragraph()
        signature.add_run(f'Rapport généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")}\n')
        signature.add_run('Conseiller en Gestion de Patrimoine\n')
        signature.add_run('Signature: _________________________')
        
        # Sauvegarder le document
        filename = f'rapport_adequation_{client.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        filepath = os.path.join(app.config['GENERATED_DOCS_FOLDER'], filename)
        doc.save(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Erreur génération rapport: {e}")
        return None

def generate_mission_letter(client):
    """Génère une lettre de mission pour le client"""
    try:
        # Créer un nouveau document
        doc = Document()
        
        # En-tête
        header = doc.add_paragraph()
        header.add_run('LETTRE DE MISSION\n').bold = True
        header.add_run('CONSEIL EN INVESTISSEMENT FINANCIER')
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        doc.add_paragraph('\n')
        
        # Date
        date_para = doc.add_paragraph()
        date_para.add_run(f'Le {datetime.now().strftime("%d/%m/%Y")}')
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph('\n')
        
        # Destinataire
        dest_para = doc.add_paragraph()
        dest_para.add_run('Madame/Monsieur ').bold = True
        dest_para.add_run(f'{client.prenom} {client.nom}\n')
        if client.adresse:
            dest_para.add_run(client.adresse)
        
        doc.add_paragraph('\n')
        
        # Objet
        objet = doc.add_paragraph()
        objet.add_run('Objet: ').bold = True
        objet.add_run('Mission de conseil en investissement financier')
        
        doc.add_paragraph('\n')
        
        # Corps de la lettre
        doc.add_paragraph('Madame, Monsieur,')
        
        doc.add_paragraph('Suite à notre entretien et à l\'analyse de votre situation patrimoniale, '
                         'nous avons le plaisir de vous proposer une mission de conseil en investissement financier.')
        
        # Profil client
        doc.add_heading('VOTRE PROFIL INVESTISSEUR', level=1)
        
        profil_text = f"""Basé sur le questionnaire complété, votre profil investisseur se caractérise par :
        
• Tolérance au risque : {client.tolerance_risque.value if client.tolerance_risque else 'Non définie'}
• Horizon d'investissement : {client.horizon_investissement.value if client.horizon_investissement else 'Non défini'}
• Score de profil : {client.profil_score if client.profil_score else 'Non défini'}/5"""
        
        doc.add_paragraph(profil_text)
        
        # Services proposés
        doc.add_heading('NOS SERVICES', level=1)
        
        services = [
            "Analyse de votre situation patrimoniale globale",
            "Recommandations d'investissement adaptées à votre profil",
            "Sélection de supports d'investissement",
            "Suivi régulier de vos investissements",
            "Reporting périodique sur la performance de vos placements"
        ]
        
        for service in services:
            doc.add_paragraph(service, style='List Bullet')
        
        # Modalités
        doc.add_heading('MODALITÉS', level=1)
        
        modalites_text = """Cette mission s'exercera dans le cadre de la réglementation en vigueur relative au conseil 
en investissement financier. Les recommandations formulées tiendront compte de votre profil investisseur, 
de vos objectifs et de votre situation financière."""
        
        doc.add_paragraph(modalites_text)
        
        # Conclusion
        doc.add_paragraph('\n')
        doc.add_paragraph('Nous espérons que cette proposition retiendra votre attention et nous tenons à votre '
                         'disposition pour tout complément d\'information.')
        
        doc.add_paragraph('Dans l\'attente de votre retour, nous vous prions d\'agréer, Madame, Monsieur, '
                         'l\'expression de nos salutations distinguées.')
        
        # Signature
        doc.add_paragraph('\n\n')
        signature = doc.add_paragraph()
        signature.add_run('Le Conseiller en Gestion de Patrimoine\n\n\n')
        signature.add_run('Signature et cachet : _________________________')
        
        # Sauvegarder le document
        filename = f'lettre_mission_{client.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        filepath = os.path.join(app.config['GENERATED_DOCS_FOLDER'], filename)
        doc.save(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Erreur génération lettre: {e}")
        return None

def generate_kyc_document(client):
    """Génère un document KYC compilant toutes les informations client"""
    try:
        doc = Document()
        
        # Titre
        title = doc.add_heading('DOCUMENT KYC - CONNAISSANCE CLIENT', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Date
        doc.add_paragraph(f'Généré le {datetime.now().strftime("%d/%m/%Y à %H:%M")}')
        
        # Informations personnelles complètes
        doc.add_heading('INFORMATIONS PERSONNELLES', level=1)
        
        info_table = doc.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'Nom et Prénom'
        info_table.cell(0, 1).text = f'{client.prenom} {client.nom}'
        
        info_table.cell(1, 0).text = 'Email'
        info_table.cell(1, 1).text = client.email
        
        info_table.cell(2, 0).text = 'Téléphone'
        info_table.cell(2, 1).text = client.telephone or 'Non renseigné'
        
        info_table.cell(3, 0).text = 'Date de naissance'
        info_table.cell(3, 1).text = client.date_naissance.strftime('%d/%m/%Y') if client.date_naissance else 'Non renseignée'
        
        info_table.cell(4, 0).text = 'Profession'
        info_table.cell(4, 1).text = client.profession or 'Non renseignée'
        
        info_table.cell(5, 0).text = 'Adresse'
        info_table.cell(5, 1).text = client.adresse or 'Non renseignée'
        
        info_table.cell(6, 0).text = 'Date d\'entrée en relation'
        info_table.cell(6, 1).text = client.date_entree_relation.strftime('%d/%m/%Y') if client.date_entree_relation else 'Non renseignée'
        
        info_table.cell(7, 0).text = 'Statut workflow'
        info_table.cell(7, 1).text = client.statut_workflow.value
        
        # Situation financière
        doc.add_heading('SITUATION FINANCIÈRE', level=1)
        
        finance_table = doc.add_table(rows=3, cols=2)
        finance_table.style = 'Table Grid'
        
        finance_table.cell(0, 0).text = 'Revenus mensuels'
        finance_table.cell(0, 1).text = f'{client.revenus_mensuels:,.0f} €' if client.revenus_mensuels else 'Non renseigné'
        
        finance_table.cell(1, 0).text = 'Charges mensuelles'
        finance_table.cell(1, 1).text = f'{client.charges_mensuelles:,.0f} €' if client.charges_mensuelles else 'Non renseigné'
        
        finance_table.cell(2, 0).text = 'Patrimoine total'
        finance_table.cell(2, 1).text = f'{client.patrimoine_total:,.0f} €' if client.patrimoine_total else 'Non renseigné'
        
        # Profil investisseur
        if client.tolerance_risque or client.horizon_investissement:
            doc.add_heading('PROFIL INVESTISSEUR', level=1)
            
            profil_table = doc.add_table(rows=5, cols=2)
            profil_table.style = 'Table Grid'
            
            profil_table.cell(0, 0).text = 'Tolérance au risque'
            profil_table.cell(0, 1).text = client.tolerance_risque.value if client.tolerance_risque else 'Non définie'
            
            profil_table.cell(1, 0).text = 'Horizon d\'investissement'
            profil_table.cell(1, 1).text = client.horizon_investissement.value if client.horizon_investissement else 'Non défini'
            
            profil_table.cell(2, 0).text = 'Expérience financière'
            profil_table.cell(2, 1).text = client.experience_financiere or 'Non renseignée'
            
            profil_table.cell(3, 0).text = 'Objectifs d\'investissement'
            profil_table.cell(3, 1).text = client.objectifs_investissement or 'Non renseignés'
            
            profil_table.cell(4, 0).text = 'Score de profil'
            profil_table.cell(4, 1).text = f'{client.profil_score}/5' if client.profil_score else 'Non calculé'
        
        # Signature
        doc.add_paragraph('\n\n')
        signature = doc.add_paragraph()
        signature.add_run('Ce document certifie la véracité des informations recueillies.\n')
        signature.add_run('Signature client : _________________________    ')
        signature.add_run('Signature conseiller : _________________________')
        
        # Sauvegarder
        filename = f'kyc_{client.id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        filepath = os.path.join(app.config['GENERATED_DOCS_FOLDER'], filename)
        doc.save(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Erreur génération KYC: {e}")
        return None

def get_recommandations(client):
    """Génère des recommandations basées sur le profil client"""
    recommandations = []
    
    if not client.tolerance_risque:
        return ["Profil de risque non défini - Veuillez compléter le questionnaire"]
    
    # Recommandations selon la tolérance au risque
    if client.tolerance_risque.name == 'FAIBLE':
        recommandations.extend([
            "Privilégier les placements sécurisés (fonds euros, obligations d'État)",
            "Limiter l'exposition aux actifs risqués à 20% maximum",
            "Diversifier sur plusieurs supports de même niveau de risque",
            "Envisager l'épargne réglementée (Livret A, LDDS) pour la liquidité"
        ])
    
    elif client.tolerance_risque.name == 'MOYENNE':
        recommandations.extend([
            "Équilibrer entre sécurité (60%) et croissance (40%)",
            "Investir en unités de compte diversifiées",
            "Privilégier les fonds mixtes équilibrés",
            "Envisager l'immobilier locatif ou les SCPI"
        ])
    
    else:  # ELEVEE
        recommandations.extend([
            "Optimiser le potentiel de croissance avec 70% d'actifs dynamiques",
            "Diversifier sur les actions européennes et internationales",
            "Envisager les investissements thématiques ou sectoriels",
            "Considérer les produits structurés avec capital non garanti"
        ])
    
    # Recommandations selon l'horizon
    if client.horizon_investissement:
        if client.horizon_investissement.name == 'COURT':
            recommandations.append("Privilégier la liquidité avec des placements facilement mobilisables")
        elif client.horizon_investissement.name == 'MOYEN':
            recommandations.append("Équilibrer rendement et accessibilité des fonds")
        else:  # LONG
            recommandations.append("Maximiser le potentiel de croissance à long terme")
    
    return recommandations
